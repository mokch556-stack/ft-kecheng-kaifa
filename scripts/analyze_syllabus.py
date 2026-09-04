# -*- coding: utf-8 -*-
"""
开课说明结构诊断脚本 — cli-anything-syllabus 配套工具
用法: python analyze_syllabus.py "课程开课说明.docx"

输出: 章节结构树 / 表格清单 / 教学目标提取 / 三级矩阵课点统计
      (K/S/A、★/☆、产出可测量关键词、周学时) / 跨课点同名主题检测
仅做自动初检，最终判断按 SKILL.md 检查清单人工复核。
"""
import os
import re
import sys
from collections import Counter

try:
    import docx
except ImportError:
    print('缺少 python-docx，请先: pip install python-docx')
    sys.exit(1)

from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# 去掉空白与控制字符，用于归一化比较
_CTRL = re.compile(r'[\x00-\x1f\x7f\u200b-\u200f\ufeff\u2028\u2029]')


def norm(s):
    s = _CTRL.sub('', s or '')
    return re.sub(r'\s+', '', s)


# 无序号中文章节标题关键词
HEADING_KW = ['课程基本信息', '课程介绍', '课程简介', '课程教学目标', '课程设计',
              '课程实施', '课程要求', '考核评价', '课程计划', '项目的创设',
              '项目一', '项目二', '项目三', '课程思政', '一级矩阵', '二级矩阵', '三级矩阵']


def is_heading(text):
    t = norm(text)
    if not t:
        return False
    # 标题应短小：超过 26 字的段落视为正文，不是标题
    if len(t) > 26:
        return False
    if re.match(r'^(一|二|三|四|五|六|七|八|九|十)、', t):
        return True
    if re.match(r'^第[一二三四五六七八九十]+[章节部分]', t):
        return True
    if t.startswith('附件') or t.startswith('附录'):
        return True
    for kw in HEADING_KW:
        if t == kw or t.startswith(kw):
            return True
    return False


MEASURE_KW = ['≥', '>=', '不少于', '提交', '链接', '截图', '二维码', '字',
              '演示', '答辩', '互评', '测试', '清单', '绘制', '撰写', '发布', '记录']

# 通用缩写黑名单（不算"平台工具名"）
TOOL_BLACKLIST = {'PPT', 'PDF', 'H5', 'API', 'URL', 'AI', 'ID', 'IP', 'CPU', 'GPU',
                  'APP', 'PC', 'UI', 'UX', 'IT', 'PS', 'EXCEL', 'WORD', 'WPS',
                  'LOGO', 'Email', 'ABCD', 'Transformer'}


def count_ksa(text):
    """粗略统计一段文本中的 K/S/A 标记与星标。"""
    k = len(re.findall(r'[KＫ]\d+', text))
    s = len(re.findall(r'[SＳ]\d+', text))
    a = len(re.findall(r'[AＡ]\d+', text))
    star = len(re.findall(r'[★☆]', text))
    return k, s, a, star


def is_third_level_matrix(table):
    """判定三级矩阵：表头第一列含'课点'，且数据行含 K/S/A 或 ★/☆。
    注意：表头可能有两行（跨页重复表头），因此数据行抽查要遍历所有行。"""
    if not table.rows:
        return False
    h0 = norm(table.rows[0].cells[0].text)
    if '课点' not in h0:
        return False
    for row in table.rows[1:]:
        sample = ' '.join(norm(c.text) for c in row.cells)
        if re.search(r'[KＳSＡA]\d+|[★☆]', sample):
            return True
    return False


def analyze(path):
    doc = docx.Document(path)
    print('=' * 70)
    print(f'文件: {os.path.basename(path)}')
    print(f'段落数: {len(doc.paragraphs)}  表格数: {len(doc.tables)}')
    print('=' * 70)

    # ---------- 1. 章节结构树 ----------
    print('\n【1】章节结构树')
    headings = [p.text.strip() for p in doc.paragraphs
                if p.text.strip() and is_heading(p.text)]
    if not headings:
        print('  (未识别到标准章节标题，可能是纯表格型文档)')
    for h in headings:
        print(f'  - {h[:40]}')
    std = ['课程基本信息', '课程介绍', '课程教学目标', '课程设计', '课程实施',
           '课程要求', '考核评价', '课程计划']
    missing = [s for s in std if not any(s in norm(h) for h in headings)]
    if missing:
        print(f'  ⚠ 疑似缺失标准章节: {missing}')

    # ---------- 2. 教学目标提取 ----------
    print('\n【2】教学目标提取')
    goals = []
    for p in doc.paragraphs:
        t = norm(p.text)
        # 匹配 "教学目标1/2/3" 或 "目标1/2/3：" 开头，排除标题与引导句
        if re.match(r'^(教学)?目标[一二三四五六七八九十1-9][:：]', t) or \
           re.match(r'^(教学)?目标[一二三四五六七八九十1-9]能', t) or \
           re.match(r'^(教学)?目标[一二三四五六七八九十1-9]：', t):
            goals.append(p.text.strip())
    if not goals:
        # 退路：含"教学目标N"的任何段落
        for p in doc.paragraphs:
            if re.search(r'(教学)?目标[一二三四五六七八九十1-9][:：能]', norm(p.text)) and len(p.text) < 200:
                goals.append(p.text.strip())
    if not goals:
        print('  (未在段落中提取到教学目标，请检查是否在表格内)')
    for g in goals[:10]:
        print(f'  - {g[:80]}')
    if goals:
        print(f'  目标条数: {len(goals)} (建议 3~5 条)')

    # ---------- 3. 表格清单与矩阵识别 ----------
    print('\n【3】表格清单')
    for ti, table in enumerate(doc.tables):
        nrow, ncol = len(table.rows), len(table.columns)
        first = norm(table.rows[0].cells[0].text)[:12] if nrow else ''
        second = norm(table.rows[1].cells[0].text)[:12] if nrow > 1 else ''
        kind = ''
        if is_third_level_matrix(table):
            kind = ' ← 三级矩阵'
        elif '课点' in first or '课程' in first:
            kind = ' ← 疑似课程/矩阵表'
        print(f'  表{ti}: {nrow}行x{ncol}列 | 首格: {first} | 次行首格: {second}{kind}')

    # ---------- 4. 三级矩阵课点统计 ----------
    print('\n【4】三级矩阵课点统计')
    week_counter = Counter()
    matrix_count = 0
    for ti, table in enumerate(doc.tables):
        if not is_third_level_matrix(table):
            continue
        matrix_count += 1
        print(f'\n  --- 表{ti} (三级矩阵) ---')
        for ri in range(1, len(table.rows)):
            cells = table.rows[ri].cells
            kp_name = norm(cells[0].text)
            if not kp_name or '达成测量' in kp_name or '学习产出' in kp_name:
                continue
            if kp_name in ('课点项目', '课点', '课点名称'):  # 跨页重复表头
                continue
            row_text = ' '.join(norm(c.text) for c in cells)
            k, s, a, star = count_ksa(row_text)
            measure_hits = [kw for kw in MEASURE_KW if kw in row_text]
            wk = norm(cells[-2].text) if len(cells) >= 7 else ''
            hr = norm(cells[-1].text) if len(cells) >= 7 else ''
            if wk.isdigit():
                week_counter[int(wk)] += 1
            flag = '✓' if len(measure_hits) >= 2 else '⚠'
            print(f'  R{ri}: {kp_name[:20]} | K{k}/S{s}/A{a} | ★☆{star} | '
                  f'可测量词{len(measure_hits)}{flag} | 周[{wk}] 学时[{hr}]')
        print('  注: 可测量词命中 ≥2 为基本达标，0~1 需人工复核')
    if matrix_count == 0:
        print('  (未识别到含"课点"的三级矩阵表格)')
    if week_counter:
        print(f'\n  周次分布: {dict(sorted(week_counter.items()))} | 上课周: {sorted(week_counter.keys())}')

    # ---------- 5. 跨课点同名主题检测 ----------
    print('\n【5】跨课点同名主题检测（启发式）')
    kp_names = []
    for table in doc.tables:
        if not is_third_level_matrix(table):
            continue
        for ri in range(1, len(table.rows)):
            n = norm(table.rows[ri].cells[0].text)
            if n and '课点' in n and '达成' not in n and '产出' not in n:
                kp_names.append(n)
    seen = {}
    for n in kp_names:
        m = re.sub(r'^[★☆]?课点\d+', '', n)
        for kw in ['图灵', '知识库', '版权', '提示词', '大模型', '智能体', '工作流',
                   '对话', '界面', '测试', '发布', '伦理', '隐私', '答辩', 'RAG', '批处理']:
            if kw in m:
                seen.setdefault(kw, []).append(n[:14])
                break
    dup = [f'  "{kw}" 出现在: {lst}' for kw, lst in seen.items() if len(lst) > 1]
    if dup:
        print('\n'.join(dup))
        print('  ⚠ 同名主题跨课点出现，请人工判断是否需要合并/分工')
    else:
        print('  (未发现明显跨课点同名主题)')

    # ---------- 6. 工具/平台名一致性初检 ----------
    print('\n【6】工具/平台名一致性初检（大写英文词）')
    tool_counter = Counter()
    for p in doc.paragraphs:
        for m in re.findall(r'[A-Z][A-Za-z]{2,}', p.text):
            if m not in TOOL_BLACKLIST:
                tool_counter[m] += 1
    for table in doc.tables:
        for row in table.rows:
            for c in row.cells:
                for m in re.findall(r'[A-Z][A-Za-z]{2,}', c.text):
                    if m not in TOOL_BLACKLIST:
                        tool_counter[m] += 1
    for name, cnt in tool_counter.most_common(10):
        flag = '⚠' if cnt < 3 else '✓'
        print(f'  {flag} {name}: {cnt} 次' + (' (仅在单处提及，检查是否漏同步)' if cnt < 3 else ''))

    print('\n' + '=' * 70)
    print('诊断完成。自动化初检仅供参考，逐项定夺请对照 SKILL.md 检查清单。')
    print('=' * 70)


def detect_vmerge(table, min_row=2):
    """定位三级矩阵数据区中的垂直合并单元格 (vMerge restart/continue)。
    min_row: 表头行数（通常为2，跨页重复表头），表头内的 vMerge 属正常结构不报告。"""
    merged = []
    seen = set()
    for ri, row in enumerate(table.rows):
        if ri < min_row:
            continue
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            if id(tc) in seen:
                continue
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                continue
            vm = tcPr.find(qn('w:vMerge'))
            if vm is not None:
                val = vm.get(qn('w:val')) or 'continue'
                merged.append((ri, ci, val))
                seen.add(id(tc))
    return merged


def structure_report(path):
    """【7】三级矩阵结构检测：垂直合并 / 达成测量行 / 学时列格式。"""
    doc = docx.Document(path)
    print('\n【7】三级矩阵结构检测')
    found_any = False
    for ti, table in enumerate(doc.tables):
        if not is_third_level_matrix(table):
            continue
        found_any = True
        print(f'\n  --- 表{ti} (三级矩阵) ---')
        # a. 垂直合并
        vm = detect_vmerge(table)
        if vm:
            groups = {}
            for ri, ci, val in vm:
                groups.setdefault(ci, []).append((ri, val))
            for ci, lst in groups.items():
                if len(lst) >= 1:
                    rows_txt = ', '.join(f'R{r}{"起点" if v=="restart" else ""}' for r, v in lst)
                    print(f'  ⚠ 第{ci}列存在垂直合并: {rows_txt} (需拆开，每课点独立填写)')
        else:
            print(f'  ✓ 无垂直合并 (vMerge)')
        # b. 达成测量行
        last = table.rows[-1]
        t0 = norm(last.cells[0].text)
        has_title = ('学习产出' in t0 and '测量' in t0)
        m_cnt = sum(1 for c in last.cells if '教学目标' in c.text and '达成测量' in c.text)
        if has_title and m_cnt >= 1:
            print(f'  ✓ 底部有教学目标达成测量行 ({m_cnt} 个目标列)')
        else:
            print(f'  ⚠ 底部缺少"学习产出及测量标准（以教学目标为单位考核）"行 (当前末行: {t0[:20]})')
        # c. 学时列格式（列头含"学时"的列）
        header = [norm(c.text) for c in table.rows[0].cells]
        for ci, h in enumerate(header):
            if '学时' in h:
                vals = [norm(table.rows[ri].cells[ci].text)
                        for ri in range(1, len(table.rows))
                        if norm(table.rows[ri].cells[ci].text)]
                has_label = any(('理论' in v or '实践' in v) for v in vals)
                if not has_label:
                    print(f'  ⚠ 学时列(第{ci}列)值为纯数字/未标注理论实践: {vals[:6]}')
                break
        # d. 产出列渠道残留（列头含"产出"）
        for ci, h in enumerate(header):
            if '产出' in h:
                channel_hits = 0
                for ri in range(1, len(table.rows)):
                    txt = table.rows[ri].cells[ci].text
                    if re.search(r'提交至|交到|上传到|发到|渠道', txt):
                        channel_hits += 1
                if channel_hits:
                    print(f'  ⚠ 产出列(第{ci}列)有 {channel_hits} 处"提交至..."渠道表述，建议去掉、渠道放正文')
                break
    if not found_any:
        print('  (未识别到含"课点"的三级矩阵表格，跳过)')
    print('\n' + '=' * 70)
    print('结构检测完成。vMerge/测量行/学时列仅自动初检，最终判断人工复核。')
    print('=' * 70)


def banned_scan(path, banned_terms):
    """【8】废弃口径反向扫描：全文（段落+表格）查找禁用词/旧口径，0 命中才算过。
    调优/改版时把本期废弃口径（旧占比/旧机制/旧叫法等）经 --banned 传入。"""
    doc = docx.Document(path)
    hits = []
    for pi, para in enumerate(doc.paragraphs):
        for w in banned_terms:
            if w and w in para.text:
                hits.append((f'段落P{pi}', w))
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen = set()
            for ci, cell in enumerate(row.cells):
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for w in banned_terms:
                    if w and w in cell.text:
                        hits.append((f'表{ti} R{ri}C{ci}', w))
    print('\n【8】废弃口径反向扫描 (banned)')
    if not banned_terms:
        print('  (未提供 --banned 词，跳过；调优/改版时把本期废弃口径列入，0 命中才算过)')
        return
    if not hits:
        print('  ✓ 0 命中，通过')
    else:
        shown = set()
        for loc, w in hits:
            key = (loc, w)
            if key in shown:
                continue
            shown.add(key)
            print(f'  ⚠ 命中"{w}" @ {loc}')
            if len(shown) >= 40:
                break
        if len(shown) < len(hits):
            print(f'  … 共 {len(hits)} 处命中，以上前 {len(shown)} 处')
        print('  ✗ 未通过：请逐处修正后再交付')


def term_group_scan(path, term_groups):
    """【9】同义词组/同物异名检测：--term-groups "A=B;C=D" 检查疑似异名并存。"""
    print('\n【9】同义词组/同物异名检测 (term-groups)')
    if not term_groups:
        print('  (未提供 --term-groups，跳过；疑似异名可指定 A=B 检测)')
        return
    doc = docx.Document(path)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                texts.append(cell.text)
    body = '\n'.join(texts)
    for g in term_groups:
        a, _, b = g.partition('=')
        if not a or not b:
            print(f'  ⚠ 无法解析分组: "{g}"（格式 A=B）')
            continue
        ca, cb = body.count(a), body.count(b)
        if ca and cb:
            print(f'  ⚠ 疑似同物异名: "{a}"×{ca} 与 "{b}"×{cb} 同时出现，请确认是否统一为一种')
        else:
            print(f'  ✓ "{a}"×{ca} / "{b}"×{cb}')


if __name__ == '__main__':
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        sys.exit(1)
    path = None
    banned = []
    groups = []
    i = 0
    while i < len(args):
        a = args[i]
        if a in ('--banned', '-b') and i + 1 < len(args):
            banned = [x.strip() for x in args[i + 1].split(',') if x.strip()]
            i += 2
        elif a in ('--term-groups', '-t') and i + 1 < len(args):
            groups = [x.strip() for x in args[i + 1].split(';') if x.strip()]
            i += 2
        else:
            path = a
            i += 1
    if not path:
        print(__doc__)
        sys.exit(1)
    analyze(path)
    structure_report(path)
    banned_scan(path, banned)
    term_group_scan(path, groups)
