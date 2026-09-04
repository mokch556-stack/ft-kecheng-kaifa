# -*- coding: utf-8 -*-
"""
docx 文本层 diff — 改版安全前置配套工具（v9.1 新增）
用法: python diff_docx.py 旧文件.docx 新文件.docx

输出: 段落+表格文本的行级差异（unified diff），供核验
      "预期改动全部命中、无意外回退（旧口径不得复活）"。
仅自动初检，最终判断人工复核。
"""
import sys
import difflib

try:
    import docx
except ImportError:
    print('缺少 python-docx，请先: pip install python-docx')
    sys.exit(1)

sys.stdout.reconfigure(encoding='utf-8')


def dump(path):
    """提取 docx 为行清单：段落 P|…，表格行 R|单元格1 || 单元格2…。"""
    doc = docx.Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append('P|' + t)
    for ti, table in enumerate(doc.tables):
        lines.append(f'=== 表{ti} ===')
        for row in table.rows:
            seen = set()
            cells = []
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                cells.append(cell.text.strip().replace('\n', '⏎'))
            lines.append('R|' + ' || '.join(cells))
    return lines


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    a, b = sys.argv[1], sys.argv[2]
    la, lb = dump(a), dump(b)
    print(f'--- {a} ({len(la)} 行) vs {b} ({len(lb)} 行) ---')
    diff = list(difflib.unified_diff(la, lb, fromfile=a, tofile=b,
                                     lineterm='', n=0))
    if not diff:
        print('✓ 文本层完全一致')
        return
    only_new = only_old = 0
    for line in diff:
        if line.startswith(('+++', '---', '@@')):
            print(line)
            continue
        print(line[:200])
        if line.startswith('+'):
            only_new += 1
        elif line.startswith('-'):
            only_old += 1
    print(f'\n统计: +{only_new} 新增/改动行, -{only_old} 删除/改动行')
    print('核验要点：预期改动应全部出现(+侧)；若出现旧口径(旧占比/旧叫法/旧机制)'
          '=> 可能被旧 Word/WPS 缓冲回退覆盖，立即停下查版本，不得继续交付。')


if __name__ == '__main__':
    main()
